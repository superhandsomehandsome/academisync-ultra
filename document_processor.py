# 多格式本地文献解析：PDF (PyMuPDF) + Word (python-docx)
import io
import os
import re
import tempfile
from collections import Counter
from typing import List, Tuple

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None


def save_temp_files(uploaded_files, max_files: int = 10) -> List[str]:
    """
    将 Streamlit 上传的文件保存到临时目录，返回本地路径列表。
    仅处理 .pdf、.docx，最多 max_files 个。
    """
    if not uploaded_files:
        return []
    paths = []
    suffix_map = {
        "application/pdf": ".pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
    }
    for f in uploaded_files[:max_files]:
        if getattr(f, "name", None) is None:
            continue
        name = (f.name or "").strip()
        if not name:
            continue
        ext = os.path.splitext(name)[1].lower()
        if ext not in (".pdf", ".docx"):
            continue
        try:
            data = f.read()
            if hasattr(f, "seek"):
                f.seek(0)
        except Exception as e:
            print(f"[警告] 读取文件 {name} 失败: {e}")
            continue
        if not data:
            print(f"[警告] 文件 {name} 内容为空，已跳过")
            continue
        suffix = suffix_map.get(getattr(f, "type", None) or "") or ext
        fd, path = tempfile.mkstemp(suffix=suffix, prefix="academisync_")
        try:
            os.write(fd, data)
        finally:
            os.close(fd)
        paths.append(path)
        print(f"[保存] 临时文件: {name} -> {path}")
    return paths


class AcademicBrain:
    """支持一次性解析最多 10 篇 PDF 或 Word，并提取实验数据、逻辑架构、专业术语供生成时优先参考。"""

    MAX_DOCS = 10

    def __init__(self):
        self.knowledge_base: List[dict] = []  # [{"source": path, "content": str}, ...]

    def load_documents(self, file_paths: List[str]) -> str:
        """从本地路径加载并解析 PDF / Word，最多 10 篇。返回状态信息。"""
        self.knowledge_base = []
        paths = [p for p in (file_paths or []) if p and os.path.isfile(p)][: self.MAX_DOCS]
        for path in paths:
            name = os.path.basename(path)
            ext = os.path.splitext(path)[1].lower()
            try:
                if ext == ".pdf":
                    content = self._parse_pdf(path)
                elif ext == ".docx":
                    content = self._parse_docx(path)
                else:
                    continue
                content = (content or "").strip()
                if content:
                    self.knowledge_base.append({"source": path, "content": content})
                    print(f"成功解析: {name}, 字数: {len(content)}")
                else:
                    print(f"解析结果为空: {name}")
            except Exception as e:
                print(f"解析失败: {name}, 错误: {e}")
        return f"已成功加载并学习 {len(self.knowledge_base)} 篇本地文献"

    def load_streamlit_files(self, uploaded_files, max_files: int = 10) -> int:
        """
        直接从 Streamlit 上传的内存流解析，避免临时文件。
        限制最多 max_files 篇，返回成功加载的篇数。
        """
        self.knowledge_base = []
        for uploaded_file in (uploaded_files or [])[:max_files]:
            file_extension = (uploaded_file.name or "").split(".")[-1].lower()
            if file_extension == "pdf":
                if fitz is None:
                    raise RuntimeError("请安装 PyMuPDF: pip install pymupdf")
                doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
                text = "".join([page.get_text() for page in doc])
                doc.close()
                self.knowledge_base.append({"title": uploaded_file.name, "content": text})
            elif file_extension == "docx":
                if DocxDocument is None:
                    raise RuntimeError("请安装 python-docx: pip install python-docx")
                doc = DocxDocument(io.BytesIO(uploaded_file.read()))
                text = "\n".join([para.text for para in doc.paragraphs])
                self.knowledge_base.append({"title": uploaded_file.name, "content": text})
        return len(self.knowledge_base)

    def load_from_uploaded_files(self, uploaded_files, max_files: int = 10) -> str:
        """
        直接从 Streamlit 上传的内存流解析，避免临时文件权限问题。
        调用 load_streamlit_files，返回状态信息供 UI 展示。
        """
        count = self.load_streamlit_files(uploaded_files, max_files)
        if count == 0:
            return "未上传文件" if not uploaded_files else "未成功解析任何文件（仅支持 PDF/DOCX）"
        return f"已成功加载并学习 {count} 篇本地文献"

    def _parse_pdf(self, path: str) -> str:
        if fitz is None:
            raise RuntimeError("请安装 PyMuPDF: pip install pymupdf")
        text = []
        with fitz.open(path) as doc:
            for page in doc:
                text.append(page.get_text() or "")
        return "\n".join(text)

    def _parse_pdf_bytes(self, data: bytes) -> str:
        """从内存字节流解析 PDF，避免临时文件。"""
        if fitz is None:
            raise RuntimeError("请安装 PyMuPDF: pip install pymupdf")
        text = []
        with fitz.open(stream=data, filetype="pdf") as doc:
            for page in doc:
                text.append(page.get_text() or "")
        return "\n".join(text)

    def _parse_docx(self, path: str) -> str:
        if DocxDocument is None:
            raise RuntimeError("请安装 python-docx: pip install python-docx")
        doc = DocxDocument(path)
        return "\n".join([(p.text or "").strip() for p in doc.paragraphs])

    def _parse_docx_bytes(self, data: bytes) -> str:
        """从内存字节流解析 Word，避免临时文件。"""
        if DocxDocument is None:
            raise RuntimeError("请安装 python-docx: pip install python-docx")
        doc = DocxDocument(io.BytesIO(data))
        return "\n".join([(p.text or "").strip() for p in doc.paragraphs])

    def _extract_sections(self, content: str) -> dict:
        """从单篇正文中粗分：实验数据、逻辑架构、专业术语。"""
        content = (content or "").strip()
        if not content:
            return {"data": "", "logic": "", "terms": ""}
        paragraphs = [p.strip() for p in re.split(r"\n\s*\n", content) if p.strip()]
        data_lines = []
        logic_lines = []
        data_re = re.compile(
            r"[%\d\.\-\+±±×÷≈≤≥P<>\s]|结果|表\s*\d|图\s*\d|实验|指标|检测|测定|P\s*[<>=]|P值|显著性"
        )
        logic_re = re.compile(
            r"首先|其次|再次|最后|因此|综上|一方面|另一方面|第一[、．]|第二[、．]|一[、．]|二[、．]|（[一二三四五六七八九十]）"
        )
        for p in paragraphs:
            if len(p) < 10:
                continue
            if data_re.search(p) and any(c.isdigit() for c in p):
                data_lines.append(p[:800])
            if logic_re.search(p):
                logic_lines.append(p[:600])
        # 专业术语：出现 2 次以上的 2–8 字中文词（简单启发）
        raw = re.sub(r"\s+", "", content)[:8000]
        words = re.findall(r"[\u4e00-\u9fff]{2,8}", raw)
        cnt = Counter(w for w in words if len(w) >= 2)
        terms = "、".join([w for w, _ in cnt.most_common(40) if _ >= 2][:25])
        return {
            "data": "\n".join(data_lines[-15:]) if data_lines else "",
            "logic": "\n".join(logic_lines[-10:]) if logic_lines else "",
            "terms": terms or "（未自动识别）",
        }

    def get_context_for_ai(self, max_per_doc: int = 2500) -> str:
        """
        生成供 AI 优先参考的本地文献精华，按「实验数据、逻辑架构、专业术语」组织。
        写作时必须优先参考这些内容。
        """
        if not self.knowledge_base:
            return ""
        parts = [
            "【以下为本地文献精华，生成文章时请优先参考其中的实验数据、逻辑架构与专业术语】"
        ]
        for item in self.knowledge_base:
            name = item.get("title") or os.path.basename(item.get("source", "")) or "未命名"
            content = (item.get("content") or "")[: max_per_doc * 3]
            sec = self._extract_sections(content)
            block = [f"\n--- 来自文献: {name} ---"]
            if sec["data"]:
                block.append("【实验数据】")
                block.append(sec["data"][:max_per_doc])
            if sec["logic"]:
                block.append("【逻辑架构】")
                block.append(sec["logic"][:max_per_doc])
            if sec["terms"]:
                block.append("【专业术语】")
                block.append(sec["terms"])
            parts.append("\n".join(block))
        return "\n".join(parts)
